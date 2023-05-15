from docx import Document
from docx.shared import Pt

from .picture import MyFloatPicture
from .section import MySection
from .table import MyTable
from .paragraph import MyParagraph
from .toc import MyToc


class MyDocument(object):
    def __init__(self, templete=None, odd_and_even_pages_header_footer=False):
        document = Document(templete)
        document.core_properties.author = ""
        document.core_properties.comments = ""
        if odd_and_even_pages_header_footer:  # 奇偶页分开
            document.settings.odd_and_even_pages_header_footer = True
        [MySection._initial_section(section) for section in document.sections]
        [MyParagraph(p).delete() for p in document.paragraphs]

        self.document = document
        self.body_section = MySection(document.sections[0])

        self.func_mapper = {
            "front_cover": self.add_cover,
            "end_cover": self.add_cover,
            "paragraph": self.add_paragraph,
            "table": self.add_table,
            "header": self.add_header,
            "header_table": self.add_header_table,
            "footer": self.add_footer,
            "footer_table": self.add_footer_table,
            "sign": self.add_sign,
            "page_break": self.add_page_break,
            "section": self.add_section,
            "table_of_contents": self.add_table_of_contents
        }

    def render(self, data, outfile):
        # data = json_to_dict(infile)
        for element in data:
            if not element:  # no data, such as: {}
                continue
            self.validate_data(element)
            element_type = element["type"]
            func = self.func_mapper[element_type]
            if element_type == "front_cover":
                func(element)
                # self.body_section = self.add_section()
                # self.body_section.linked_to_previous(value=False)
            elif element_type == "end_cover":
                # section = self.add_section()
                # section.linked_to_previous(value=False)
                func(element)
            elif element_type == "section":
                section = self.add_section()
                section.linked_to_previous(value=False)
            else:
                func(element)
        self.document.save(outfile)

    def validate_data(self, element):
        assert isinstance(element, dict), "expected dict, but get %s" %element

    def add_cover(self, element):
        src = element["value"]
        run = self.document.add_paragraph().add_run()
        style = element["style"]
        if not style:
            style = {}
        width = style.get("width", Pt(595))
        height = style.get("height", Pt(842))
        text_wrap = style.get("text_wrap", "wrapTopAndBottom")
        pic = MyFloatPicture.create(
            run=run, src=src, width=width, height=height,
            pos_x=0, pos_y=0, anchor_x="page", anchor_y="page",
            text_wrap=text_wrap)
        return pic

    def add_paragraph(self, element):
        text = element["value"]
        style = element["style"]
        if not style:
            style = {}
        paragraph_style = style.get("paragraph", None)
        font_style = style.get("font", None)
        return MyParagraph.create_from_xml_like_text(
            text,
            self.document,
            paragraph_style=paragraph_style,
            font_style=font_style)

    def add_table(self, element):
        data = element["value"]
        style = element["style"]
        if not style:
            style = []
        return MyTable.create_from_2d_array(
            data, self.document,
            width=self.document._block_width,
            style=style
        )

    def add_header(self, element):
        data = element["value"]
        style = element["style"]
        if not style:
            style = {}
        odd_font_style = style.get("odd_font", None)
        odd_paragraph_style = style.get("odd_paragraph", None)
        even_font_style = style.get("even_font", None)
        even_paragraph_style = style.get("even_paragraph", None)
        self.body_section.create_header_from_xml_like_text(
            data,
            odd_font_style=odd_font_style, odd_paragraph_style=odd_paragraph_style,
            even_font_style=even_font_style, even_paragraph_style=even_paragraph_style
        )

    def add_header_table(self, element):
        data = element["value"]
        style = element["style"]
        if not style:
            style = {}
        odd_style = style.get("odd_style", None)
        even_style = style.get("even_style", None)
        self.body_section.create_header_from_2d_array(
            data,
            odd_style=odd_style,
            even_style=even_style
        )

    def add_footer(self, element):
        data = element["value"]
        style = element["style"]
        if not style:
            style = {}
        odd_font_style = style.get("odd_font", None)
        odd_paragraph_style = style.get("odd_paragraph", None)
        even_font_style = style.get("even_font", None)
        even_paragraph_style = style.get("even_paragraph", None)
        self.body_section.create_footer_from_xml_like_text(
            data,
            odd_font_style=odd_font_style, odd_paragraph_style=odd_paragraph_style,
            even_font_style=even_font_style, even_paragraph_style=even_paragraph_style
        )

    def add_footer_table(self, element):
        data = element["value"]
        style = element["style"]
        if not style:
            style = {}
        odd_style = style.get("odd_style", None)
        even_style = style.get("even_style", None)
        self.body_section.create_footer_from_2d_array(
            data,
            odd_style=odd_style,
            even_style=even_style
        )

    def add_sign(self, element):
        data = element["value"]
        style = element["style"]
        if not style:
            style = {}
        sign = data["sign"]   # 签名
        stamp = data["stamp"] # 章
        stamp_attr = data.get("stamp_attr", {})
        table = MyTable.create_from_2d_array(
            sign, self.document,
            width=self.document._block_width,
            style=style
        )
        if stamp:
            r = table.table.cell(0, 0).paragraphs[0].add_run()
            width = stamp_attr.get("width", Pt(100))
            height = stamp_attr.get("height", None)
            pos_x = stamp_attr.get("pos_x", Pt(280))
            pos_y = stamp_attr.get("pos_y", Pt(-28))
            MyFloatPicture.create(
                stamp, r,
                width=width, height=height,
                pos_x=pos_x, pos_y=pos_y,
                anchor_x="character", anchor_y="paragraph",
                text_wrap="wrapNone"
            )
        return table

    def add_page_break(self, element=None):
        return self.document.add_page_break()

    def add_section(self, element=None):
        return MySection.create(self.document)

    def add_table_of_contents(self, element):
        data = element["value"]
        if not data:
            data = "TOC \\o \"1-4\" \\h \\z \\u"
        return MyToc.create(data, self.document)

    def save(self, filename):
        self.document.save(filename)