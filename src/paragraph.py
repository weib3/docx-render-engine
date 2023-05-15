from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.blkcntnr import BlockItemContainer
from docx.shared import RGBColor, Pt

from .utils import bool_str, parse_from_xml_like_text, hex_to_rgb, \
    point_unit_guess, str_to_numeric
from .picture import MyInlinePicture


class MyParagraph(object):
    def __init__(self, paragraph: Paragraph):
        self.paragraph = paragraph

    @classmethod
    def create_from_xml_like_text(cls, text:str, parent: BlockItemContainer,
                                  paragraph_style=None, font_style=None):
        if not paragraph_style:
            paragraph_style = {}
        if not font_style:
            font_style = {}
        return_list = []
        text = str(text)
        for i in text.split("\n"):
            p = parent.add_paragraph()
            instance = cls(p)
            instance.set_paragraph_format(**paragraph_style)
            instance.add_runs(i, font_style)
            return_list.append(instance)
        return return_list

    def add_runs(self, text: str, style=None):
        if not style:
            style = {}
        parsed = parse_from_xml_like_text(text)
        for i in parsed:
            child_style = i["style"]
            if style:
                [child_style.setdefault(_, style[_]) for _ in style]
            if i["tag"] == "pic":
                src = child_style.pop("src")
                MyRun.create_from_inline_picture(src, self.paragraph, child_style)
            elif i["tag"].lower() in ["page", "numpages", "sectionpages"]:
                MyRun.create_from_field(i["tag"].upper(), self.paragraph, child_style)
            else:
                if i["text"]:  #有文本的时候才会添加
                    MyRun.create_from_text(i["text"], self.paragraph, child_style)

    def set_font(self, **kwargs):
        for run in self.paragraph.runs:
            MyRun(run).set_font(**kwargs)

    def set_paragraph_format(self, **kwargs):
        pfmt = self.paragraph.paragraph_format
        if "alignment" in kwargs:
            pfmt.alignment = str_to_numeric(kwargs["alignment"])
        if "first_line_indent" in kwargs:
            pfmt.first_line_indent = str_to_numeric(kwargs["first_line_indent"])
        if "keep_together" in kwargs:
            pfmt.keep_together = bool_str(kwargs["keep_together"])
        if "keep_with_next" in kwargs:
            pfmt.keep_with_next = bool_str(kwargs["keep_with_next"])
        if "line_spacing" in kwargs:
            pfmt.line_spacing = kwargs["line_spacing"]
        if "page_break_before" in kwargs:
            pfmt.page_break_before = bool_str(kwargs["page_break_before"])
        if "space_after" in kwargs:
            pfmt.space_after = str_to_numeric(kwargs["space_after"])
        if "space_before" in kwargs:
            pfmt.space_before = str_to_numeric(kwargs["space_before"])
        if "widow_control" in kwargs:
            pfmt.widow_control = bool_str(kwargs["widow_control"])
        if "outline_lvl" in kwargs:
            self.set_outline_lvl(kwargs["outline_lvl"])
        if "border" in kwargs:
            self.set_border(**kwargs["border"])

    def set_outline_lvl(self, outlinelvl):
        """
        Set paragraph outline level. Useful for generate table of contents.
        """
        outlinelvl = str(outlinelvl)
        pPr = self.paragraph.paragraph_format.element.get_or_add_pPr()
        outline = pPr.first_child_found_in("w:outlineLvl")
        if outline is None:
            element = OxmlElement("w:outlineLvl")
            element.set(qn('w:val'), outlinelvl)
            pPr.append(element)
        else:
            outline.set(qn('w:val'), outlinelvl)

    def add_horizontalline(self):
        self.set_border(
            bottom={"sz": 6, "val": "single", "color": "#000000"}
        )

    def add_border(self):
        self.set_border(
            top={"sz": 6, "val": "single", "color": "#000000"},
            bottom={"sz": 6, "val": "single", "color": "#000000"},
            start={"sz": 6, "val": "single", "color": "#000000"},
            end={"sz": 6, "val": "single", "color": "#000000"}
        )

    def set_border(self, **kwargs):
        p = self.paragraph._p
        pPr = p.get_or_add_pPr()

        # check for tag existnace, if none found, then create one
        pBdr = pPr.first_child_found_in("w:pBdr")
        if pBdr is None:
            pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(
            pBdr,
            'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku',
            'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct',
            'w:autoSpaceDE', 'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd',
            'w:snapToGrid', 'w:spacing', 'w:ind', 'w:contextualSpacing',
            'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
            'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
            'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
            'w:pPrChange')

        # left and right are used in older version word and all version of wps.
        if "start" in kwargs:
            kwargs["left"] = kwargs["start"]
        if "end" in kwargs:
            kwargs["right"] = kwargs["end"]

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV',
                     'left', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = pBdr.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    pBdr.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def delete(self):
        p = self.paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


class MyRun(object):
    def __init__(self, run):
        self.run = run

    @classmethod
    def create_from_text(cls, text: str, parent: Paragraph, style=None):
        if not style:
            style = {}
        r = parent.add_run(text)
        instance = cls(r)
        instance.set_font(**style)
        return instance

    @classmethod
    def create_from_field(cls, instrText_text, parent: Paragraph, style=None):
        if not style:
            style = {}
        r = parent.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = instrText_text

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r._r.append(fldChar1)
        r._r.append(instrText)
        r._r.append(fldChar2)
        instance = cls(r)
        instance.set_font(**style)
        return instance

    @classmethod
    def create_from_inline_picture(cls, src, parent: Paragraph, style=None):
        if not style:
            style = {}
        r = parent.add_run()
        MyInlinePicture.create(src, r, **style)
        return cls(r)

    def set_font(self, **kwargs):
        font = self.run.font
        if "bold" in kwargs:
            font.bold = bool_str(kwargs["bold"])
        if "color" in kwargs:
            color = hex_to_rgb(kwargs["color"])
            font.color.rgb = RGBColor(*color)
        if "italic" in kwargs:
            font.italic = bool_str(kwargs["italic"])
        if "name" in kwargs:
            font.name = kwargs["name"]
            self.run._element.rPr.rFonts.set(qn('w:eastAsia'),
                                        kwargs["name"])
        if "shadow" in kwargs:
            font.shadow = bool_str(kwargs["shadow"])
        if "size" in kwargs:
            font.size = point_unit_guess(str_to_numeric(kwargs["size"]))
        if "subscript" in kwargs:
            font.subscript = bool_str(kwargs["subscript"])
        if "superscript" in kwargs:
            font.superscript = bool_str(kwargs["superscript"])
        if "underline" in kwargs:
            font.underline = bool_str(kwargs["underline"])

