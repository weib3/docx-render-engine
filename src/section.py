from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .table import MyTable
from .paragraph import MyParagraph
from .utils import bool_str

class MySection(object):
    def __init__(self, section):
        self.section = section

    @classmethod
    def create(cls, document, restart_pagenum=True):
        section = document.add_section()
        section = cls._initial_section(section)
        if restart_pagenum:
            element = OxmlElement("w:pgNumType")
            element.set(qn("w:start"), "1")
            section._sectPr.append(element)

        return cls(section)

    @staticmethod
    def _initial_section(section):
        section.page_height = Pt(842)
        section.page_width = Pt(595)
        [MyParagraph(p).delete() for p in section.footer.paragraphs if
         not p.runs]
        [MyParagraph(p).delete() for p in section.even_page_footer.paragraphs if
         not p.runs]
        [MyParagraph(p).delete() for p in section.header.paragraphs if
         not p.runs]
        [MyParagraph(p).delete() for p in section.even_page_header.paragraphs if
         not p.runs]
        return section

    def create_footer_from_xml_like_text(
            self, text,
            odd_font_style=None, odd_paragraph_style=None,
            even_font_style=None, even_paragraph_style=None):

        MyParagraph.create_from_xml_like_text(
            text, self.section.footer,
            font_style=odd_font_style,
            paragraph_style=odd_paragraph_style
        )
        MyParagraph.create_from_xml_like_text(
            text, self.section.even_page_footer,
            font_style=even_font_style,
            paragraph_style=even_paragraph_style
        )

    def create_footer_from_2d_array(self, data, odd_style=None, even_style=None):
        MyTable.create_from_2d_array(
            data,
            self.section.footer,
            width=self._width,
            style=odd_style
        )
        MyTable.create_from_2d_array(
            data,
            self.section.even_page_footer,
            width=self._width,
            style=even_style
        )

    def create_header_from_xml_like_text(
            self, text,
            odd_font_style=None, odd_paragraph_style=None,
            even_font_style=None, even_paragraph_style=None):

        MyParagraph.create_from_xml_like_text(
            text, self.section.header,
            font_style=odd_font_style,
            paragraph_style=odd_paragraph_style
        )
        MyParagraph.create_from_xml_like_text(
            text, self.section.even_page_header,
            font_style=even_font_style,
            paragraph_style=even_paragraph_style
        )

    def create_header_from_2d_array(self, data, odd_style=None, even_style=None):
        MyTable.create_from_2d_array(
            data,
            self.section.header,
            width=self._width,
            style=odd_style
        )
        MyTable.create_from_2d_array(
            data,
            self.section.even_page_header,
            width=self._width,
            style=even_style
        )

    def linked_to_previous(self, value):
        value = bool_str(value)
        self.section.header.is_linked_to_previous = value
        self.section.even_page_header.is_linked_to_previous = value
        self.section.footer.is_linked_to_previous = value
        self.section.even_page_footer.is_linked_to_previous = value

    @property
    def _width(self):
        return self.section._document_part.document._block_width
