from typing import Union

from docx.document import Document
from docx.blkcntnr import BlockItemContainer
from docx.table import Table
from docx.table import _Row, _Column, _Cell
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from .utils import bool_str, is_digit, str_to_numeric
from .paragraph import MyParagraph


class MyTable(object):
    def __init__(self, table: Table):
        self.table = table

    @classmethod
    def create_from_2d_array(cls, data,
                             parent: Union[Document, BlockItemContainer],
                             width,
                             style=None):
        if not style:
            style = []
        cls._data_format_check(data)
        if isinstance(parent, Document):
            table = parent.add_table(rows=len(data), cols=len(data[0]))
        else:
            table = parent.add_table(rows=len(data), cols=len(data[0]), width=width)
        instance = MyTable(table)
        for row_idx, row in enumerate(data):
            for col_idx, text in enumerate(row):
                MyCell(table.cell(row_idx, col_idx)).set_text(text)

        # 设置样式。
        # 不能先合并单元格，否则合并单元格无法选中设置边框等样式
        for i in style:
            instance.set_style_by_selectors(i["selectors"], i["style"])

        # 合并单元格
        merged_cells = cls._get_merged_cells(data)
        for i in merged_cells:
            cell1 = table.cell(*i[0])
            cell2 = table.cell(*i[1])
            MyCell(cell1).merge_another_cell(cell2)
        return instance

    def set_style_by_selectors(self, selectors, style):
        def select_list(idx, inlist) -> list:
            if is_digit(idx):
                idx = int(idx)
                if idx > 0:
                    return [inlist[idx - 1]]
                elif idx < 0:
                    return [inlist[idx]]
                else:
                    raise Exception("idx can not be 0.")
            elif idx == "odd":  # 奇数
                return inlist[::2]
            elif idx == "even":  # 偶数
                return inlist[1::2]
            else:
                raise Exception(
                    "idx must be int, odd or even. %s provided." % idx)

        table = self.table
        for s in selectors:
            if s.startswith("table"):
                self.set_style(style)
            elif s.startswith("row"):
                idx = s.split(".")[1]
                for row in select_list(idx, list(table.rows)):
                    MyRow(row).set_style(style)
            elif s.startswith("column"):
                idx = s.split(".")[1]
                for column in select_list(idx, list(table.columns)):
                    MyColumn(column).set_style(style)
            elif s.startswith("cell"):
                row_idx = s.split(".")[1]
                col_idx = s.split(".")[2]
                for row in select_list(row_idx, list(table.rows)):
                    for cell in select_list(col_idx, list(row.cells)):
                        MyCell(cell).set_style(style)
            else:
                return

    def set_style(self, style):
        for cell in self.table._cells:
            MyCell(cell).set_style(style)
        if "autofit" in style:
            self.table.autofit = bool_str(style["autofit"])
        if "alignment" in style:
            self.set_alignment(style["alignment"])
        if "width" in style:
            self.set_width(style["width"])
        if "first_row_repeat_in_each_page" in style:
            self.set_first_row_repeat_in_each_page(style["first_row_repeat_in_each_page"])
        if "row_height" in style:
            self.set_row_height(style["row_height"])
        if "column_width" in style:
            self.set_column_width(style["column_width"])

    def set_font(self, **kwargs):
        for row in self.table.rows:
            MyRow(row).set_font(**kwargs)

    def set_paragraph_format(self, **kwargs):
        for row in self.table.rows:
            MyRow(row).set_paragraph_format(**kwargs)

    def set_border(self, **kwargs):
        for row in self.table.rows:
            MyRow(row).set_border(**kwargs)

    def set_margin(self, **kwargs):
        for row in self.table.rows:
            MyRow(row).set_margin(**kwargs)

    def set_shading(self, color):
        for row in self.table.rows:
            MyRow(row).set_shading(color)

    def clean_border(self):
        for row in self.table.rows:
            MyRow(row).clean_border()

    def add_border(self):
        for row in self.table.rows:
            MyRow(row).add_border()

    def set_alignment(self, alignment):
        self.table.alignment = str_to_numeric(alignment)

    def set_width(self, width):
        element = self.table._tblPr.find(qn("w:tblW"))
        element.set(qn("w:type"), "dxa")
        element.set(qn("w:w"), str(width))

    def set_first_row_repeat_in_each_page(self, first_row_repeat_in_each_page):
        trPr = self.table.rows[0]._tr.get_or_add_trPr()
        element = trPr.find(qn("w:tblHeader"))
        if bool_str(first_row_repeat_in_each_page):
            if element is None:
                element = OxmlElement("w:tblHeader")
                trPr.append(element)
        else:
            if element is not None:
                trPr.remove(element)

    def set_vertical_alignment(self, alignment):
        for row in self.table.rows:
            MyRow(row).set_vertical_alignment(alignment)

    def set_row_height(self, height):
        for row in self.table.rows:
            MyRow(row).set_height(height)

    def set_column_width(self, width):
        for column in self.table.columns:
            MyColumn(column).set_width(width)

    def delete(self):
        table = self.table._element
        table.getparent().remove(table)
        table._tbl = table._element = None

    @staticmethod
    def _data_format_check(data):
        # 每行数据元素数相同
        assert len(set([len(i) for i in data])) == 1, "%s have no equal " \
                                                      "length of cols" %data

    @staticmethod
    def _get_merged_cells(data):
        def get_rowspan_cells(data):
            merge_cells = []
            last_col = 1
            for row_idx, row in enumerate(data):
                for col_idx, cell in enumerate(row):
                    if cell == "~~":
                        merge_cells.append(
                            [(row_idx, last_col), (row_idx, col_idx)])
                    else:
                        last_col = col_idx
            return merge_cells

        def get_colspan_cells(data):
            merge_cells = []
            last_row_list = [1 for i in data[0]]
            for row_idx, row in enumerate(data):
                for col_idx, cell in enumerate(row):
                    if cell == "^^":
                        last_row = last_row_list[col_idx]
                        merge_cells.append(
                            [(last_row, col_idx), (row_idx, col_idx)])
                    else:
                        last_row_list[col_idx] = row_idx
            return merge_cells

        row_span = get_rowspan_cells(data)
        col_span = get_colspan_cells(data)

        row_span_dict = {i[1]: i[0] for i in row_span}
        for i in col_span:
            if i[0] in row_span_dict:
                i[0] = row_span_dict[i[0]]
        return [*row_span, *col_span]


class _RowColumnBase(object):
    def __init__(self, object: Union[_Row, _Column]):
        self.object = object

    def set_font(self, **kwargs):
        for cell in self.object.cells:
            MyCell(cell).set_font(**kwargs)

    def set_paragraph_format(self, **kwargs):
        for cell in self.object.cells:
            MyCell(cell).set_paragraph_format(**kwargs)

    def set_border(self, **kwargs):
        for cell in self.object.cells:
            MyCell(cell).set_border(**kwargs)

    def set_margin(self, **kwargs):
        for cell in self.object.cells:
            MyCell(cell).set_margin(**kwargs)

    def set_shading(self, color):
        for cell in self.object.cells:
            MyCell(cell).set_shading(color)

    def clean_border(self):
        for cell in self.object.cells:
            MyCell(cell).clean_border()

    def add_border(self):
        for cell in self.object.cells:
            MyCell(cell).add_border()

    def set_vertical_alignment(self, alignment):
        for cell in self.object.cells:
            MyCell(cell).set_vertical_alignment(alignment)

    def merge_all_cells(self):
        cell1 = self.object.cells[0]
        for i in range(1, len(self.object.cells)):
            cell2 = self.object.cells[i]
            cell1.merge(cell2)


class MyRow(_RowColumnBase):
    def __init__(self, object: _Row):
        super().__init__(object)
        self.row = self.object

    def insert_row_after(self):
        tr = self.row._tr
        new_tr = OxmlElement("w:tr")
        tr.addnext(new_tr)
        for cell in self.row.cells:
            tc = new_tr.add_tc()
            # tc.width = cell.width
        return _Row(new_tr, self.row.table)

    def insert_row_before(self):
        tr = self.row._tr
        new_tr = OxmlElement("w:tr")
        tr.addprevious(new_tr)
        for cell in self.row.cells:
            tc = new_tr.add_tc()
            # tc.width = cell.width
        return _Row(new_tr, self.row.table)

    def set_height(self, height):
        self.row.height = str_to_numeric(height)

    def set_style(self, style):
        for cell in self.object.cells:
            MyCell(cell).set_style(style)
        # cell can not set height
        if "height" in style:
            self.set_height(str_to_numeric(style["height"]))


class MyColumn(_RowColumnBase):
    def __init__(self, object: _Column):
        super().__init__(object)
        self.column = self.object

    def set_width(self, width):
        for cell in self.object.cells:
            MyCell(cell).set_width(width)
        # self.column.width = str_to_numeric(width)

    def set_style(self, style):
        for cell in self.object.cells:
            MyCell(cell).set_style(style)
        # cell can set width
        # if "width" in style:
        #     self.set_width(str_to_numeric(style["width"]))


class MyCell(object):
    def __init__(self, cell: _Cell):
        self.cell = cell

    def set_style(self, style):
        if "paragraph" in style:
            self.set_paragraph_format(**style["paragraph"])
        if "font" in style:
            self.set_font(**style["font"])
        if "border" in style:
            self.set_border(**style["border"])
        if "shading" in style:
            self.set_shading(style["shading"])
        if "vertical_alignment" in style:
            self.set_vertical_alignment(style["vertical_alignment"])
        if "margin" in style:
            self.set_margin(**style["margin"])
        if "width" in style:
            self.set_width(style["width"])

    def set_text(self, text):
        [MyParagraph(p).delete() for p in self.cell.paragraphs]
        MyParagraph.create_from_xml_like_text(text, self.cell)

    def set_font(self, **kwargs):
        """
        Set cell font style.

        Usage:
            set_font_style(
                name="方正兰亭细黑_GBK",
                bold=True,
                size=Pt(10),
                italic=False,
                color="#FFFFFF"
            )
        """
        for p in self.cell.paragraphs:
            MyParagraph(p).set_font(**kwargs)

    def set_paragraph_format(self, **kwargs):
        """
        Set cell paragraph format style.

        Args:
            **kwargs: key-values for paragraph format.

        Usage:
            set_paragraph_format(
                space_before=0,
                space_after=0
            )
        """
        for p in self.cell.paragraphs:
            MyParagraph(p).set_paragraph_format(**kwargs)

    def set_border(self, **kwargs):
        """
        Set cell's border style. Copied from stackoverflow.
        Border style see http://officeopenxml.com/WPtableBorders.php.

        Args:
            **kwargs: Dict which key is edges and values is attribute.
                      Edges include: start, top, end, bottom, insideH, insideV.

        Usage:
            set_border(
                top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
                bottom={"sz": 12, "color": "#00FF00", "val": "single"},
                start={"sz": 24, "val": "dashed", "shadow": "true"},
                end={"sz": 12, "val": "dashed"}
            )
        """
        tc = self.cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tc_borders = tcPr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tcPr.append(tc_borders)

        # left and right are used in older version word and all version of wps.
        if "start" in kwargs:
            kwargs["left"] = kwargs["start"]
        if "end" in kwargs:
            kwargs["right"] = kwargs["end"]

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV',
                     'left', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tc_borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def set_margin(self, **kwargs):
        """
        Set cell's margin.
        Margin style see http://officeopenxml.com/WPtableCellProperties-Margins.php.

        Args:
            **kwargs: Dict which key is edges and values is attribute.
                      Edges include: start, top, end, bottom.

        Usage:
            set_margin(
                top=50,
                bottom=50,
                start=50,
                end=50
            )
        """
        tc = self.cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tc_margin = tcPr.first_child_found_in('w:tcMar')
        if tc_margin is None:
            tc_margin = OxmlElement('w:tcMar')
            tcPr.append(tc_margin)

        # left and right are used in older version word and all version of wps.
        if "start" in kwargs:
            kwargs["left"] = kwargs["start"]
        if "end" in kwargs:
            kwargs["right"] = kwargs["end"]

        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'left', 'right'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tc_margin.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tc_margin.append(element)

                # set data
                element.set(qn('w:w'), str(edge_data))
                element.set(qn('w:type'), 'dxa')

    def set_shading(self, color):
        """
        Set cell shading color.

        Usage:
            set_shading(color="#1F5C8B")
        """
        tc = self.cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shading = parse_xml(
            r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        tc_pr.append(shading)

    def clean_border(self):
        """
        Clean cell border.
        """
        self.set_border(
            start={"val": "none", "sz": "1", "space": "0"},
            end={"val": "none", "sz": "1", "space": "0"},
            top={"val": "none", "sz": "1", "space": "0"},
            bottom={"val": "none", "sz": "1", "space": "0"}
        )

    def set_vertical_alignment(self, alignment):
        self.cell.vertical_alignment = str_to_numeric(alignment)

    def add_border(self):
        self.set_border(
            top={"sz": 6, "val": "single", "color": "#000000"},
            bottom={"sz": 6, "val": "single", "color": "#000000"},
            start={"sz": 6, "val": "single", "color": "#000000"},
            end={"sz": 6, "val": "single", "color": "#000000"}
        )

    def set_width(self, width):
        self.cell.width = str_to_numeric(width)

    def merge_another_cell(self, cell2):
        """
        Merge two cells with remove cell2's text.
        """
        for p in cell2.paragraphs:
            p.clear()
        self.cell.merge(cell2)

    def add_table(self, data, style=None):
        return MyTable.create_from_2d_array(
            data, self.cell, width=self._width, style=style
        )

    @property
    def _width(self):
        return self.cell.width